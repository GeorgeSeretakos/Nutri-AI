from docx import Document
from nutrition_ai.models.diet_plan import DietPlan
from pathlib import Path


class DocxExporter:

    @staticmethod
    def export_diet_to_docx(diet: DietPlan, output_path: str):
        """
        Export a DietPlan object to a DOCX document using the expert template style.
        """

        doc = Document()

        # Header
        doc.add_heading(f"ΟΝΟΜΑ: {diet.user_name}", level=1)
        doc.add_paragraph(f"ΗΜΕΡΟΜΗΝΙΑ ΓΕΝΝΗΣΗΣ: {diet.birth_date}")
        doc.add_paragraph(f"ΥΨΟΣ: {diet.height}")
        doc.add_paragraph(f"ΑΣΚΗΣΗ: {diet.exercise}")
        doc.add_paragraph(f"ΣΤΟΧΟΣ: {diet.goal}")
        doc.add_paragraph(" ")

        # Sections
        for key, section in diet.sections.items():
            doc.add_heading(section.title, level=2)
            for opt in section.options:
                doc.add_paragraph(f"- {opt.description}")
                if opt.frequency:
                    doc.add_paragraph(f"  Συχνότητα: {opt.frequency}")
                if opt.recommended_products:
                    doc.add_paragraph(f"  Προτεινόμενο προϊόν: {opt.recommended_products}")
                if opt.comments:
                    doc.add_paragraph(f"  Σχόλια: {opt.comments}")
                doc.add_paragraph(" ")

        # Fruit List
        doc.add_heading("Λίστα Φρούτων", level=2)
        for fruit in diet.fruit_list:
            doc.add_paragraph(f"- {fruit.name}: {fruit.quantity}")

        # Additional Notes
        if diet.additional_notes:
            doc.add_heading("Σχόλια", level=2)
            doc.add_paragraph(diet.additional_notes)

        # Save
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)
