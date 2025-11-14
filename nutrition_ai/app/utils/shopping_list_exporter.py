from docx import Document
from docx.shared import Pt
from pathlib import Path
from nutrition_ai.agents.shopping_list_agent import ShoppingListAgent


class ShoppingListExporter:

    @staticmethod
    def export_to_docx(shopping_list: ShoppingListAgent, output_path: str):
        """
        Export shopping list to a DOCX file with categories and checkboxes.
        """

        doc = Document()
        doc.add_heading("Λίστα Αγορών Εβδομάδας", level=1)

        checkbox = "☐"  # Unicode checkbox

        for category, items in shopping_list.items.items():
            if not items:
                continue

            # Category heading
            h = doc.add_heading(category.upper(), level=2)
            h.runs[0].font.size = Pt(14)

            for ingredient, qty in items.items():
                p = doc.add_paragraph()
                run = p.add_run(f"{checkbox} {ingredient} — {qty} φορές")
                run.font.size = Pt(12)

        # Save file
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm


class ShoppingListExporterPDF:

    @staticmethod
    def export_to_pdf(shopping_list: ShoppingListAgent, output_path: str):
        """
        Export shopping list to a clean PDF with categories and checkboxes.
        """

        doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        checkbox = "☐"

        title = Paragraph("<b>Λίστα Αγορών Εβδομάδας</b>", styles["Title"])
        story.append(title)
        story.append(Spacer(1, 0.5 * cm))

        for category, items in shopping_list.items.items():
            if not items:
                continue

            # Category header
            cat_header = Paragraph(f"<b>{category.upper()}</b>", styles["Heading2"])
            story.append(cat_header)
            story.append(Spacer(1, 0.2 * cm))

            for ingredient, qty in items.items():
                line = Paragraph(f"{checkbox} {ingredient} — {qty} φορές", styles["BodyText"])
                story.append(line)
                story.append(Spacer(1, 0.1 * cm))

            story.append(Spacer(1, 0.3 * cm))

        doc.build(story)
